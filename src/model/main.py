#!/usr/bin/env python3
import pandas as pd
import iris
from sentence_transformers import SentenceTransformer
from langchain.agents import create_agent
from langchain.agents.middleware import SummarizationMiddleware
from langgraph.checkpoint.memory import InMemorySaver
from openai import OpenAI

from langchain_openai import ChatOpenAI
from langchain.agents import create_agent
from langchain.agents.middleware import SummarizationMiddleware
from langgraph.checkpoint.memory import InMemorySaver
from langchain_core.tools import tool
from docx import Document
from dotenv import load_dotenv

import difflib
import os

load_dotenv()

api_key = os.getenv("API_KEY")

from openai import OpenAI

client = OpenAI(
  api_key=api_key
)
OPENAI_API_KEY = api_key

class RAGChatbot:
    def __init__(self):

        # Directory with .docx templates
        self.doc_root = os.getenv("DATA_PATH")

        self.current_doc: Document | None = None
        self.current_doc_path: str | None = None
        self.current_doc_name: str | None = None

        conn = iris.connect("localhost", 32782, "DEMO", "_SYSTEM", "ISCDEMO")
        self.cursor = conn.cursor()

        self.embedding_model = self.get_embedding_model()
        self.agent = self.create_chatbot()
        self.config = {"configurable": {"thread_id": "1"}}

    # ------------- MODELS / VECTOR SEARCH ------------- #

    def get_embedding_model(self):
        return SentenceTransformer("all-MiniLM-L6-v2")

    def vector_search(self, user_prompt: str):
        search_vector = self.embedding_model.encode(
            user_prompt,
            normalize_embeddings=False,
            show_progress_bar=False,
        ).tolist()

        search_sql = """
            SELECT TOP 5 filename, content
            FROM VectorSearch.ORGstruct
            ORDER BY VECTOR_COSINE(vector, TO_VECTOR(?,DOUBLE)) DESC
        """
        self.cursor.execute(search_sql, [str(search_vector)])
        results = self.cursor.fetchall()
        return [f"Text z dokumentu {x} -> {y}" for x, y in results]

    # ------------- WORD DOCUMENT HELPERS ------------- #

    @staticmethod
    def _para_text(paragraph) -> str:
        """Read paragraph text via run XML (avoids lxml XPath bug)."""
        pieces: list[str] = []
        for run in paragraph.runs:
            r = run._r
            for child in r:
                tag = child.tag
                if tag.endswith("}t") and child.text:
                    pieces.append(child.text)
                elif tag.endswith("}tab"):
                    pieces.append("\t")
                elif tag.endswith(("}br", "}cr")):
                    pieces.append("\n")
        return "".join(pieces)

    def _set_para_text(self, paragraph, new_text: str) -> None:
        """Replace paragraph content with a single run containing new_text."""
        p = paragraph._p
        for child in list(p):
            p.remove(child)

        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_text
        r.append(t)
        p.append(r)

    def _cell_text(self, cell) -> str:
        """Concatenate visible text of all paragraphs in a cell."""
        return "\n".join(self._para_text(p) for p in cell.paragraphs).strip()

    def _set_cell_text(self, cell, new_text: str) -> None:
        """Overwrite a cell with a single paragraph containing new_text."""
        if not cell.paragraphs:
            p = cell.add_paragraph("")
            self._set_para_text(p, new_text)
            return

        for i, p in enumerate(cell.paragraphs):
            if i == 0:
                self._set_para_text(p, new_text)
            else:
                self._set_para_text(p, "")

    def _doc_to_text(self, doc: Document, max_chars: int = 4000) -> str:
        """Plain text view of a document, truncated if needed."""
        parts = [self._para_text(p).strip() for p in doc.paragraphs]
        parts = [p for p in parts if p]
        text = "\n".join(parts)
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[Text zkrácen kvůli délce…]"
        return text

    def _list_docx_files(self):
        if not os.path.isdir(self.doc_root):
            return []
        return [
            f
            for f in os.listdir(self.doc_root)
            if f.lower().endswith(".docx")
            and os.path.isfile(os.path.join(self.doc_root, f))
        ]

    @staticmethod
    def _make_working_filename(filename: str) -> str:
        """Generate unique filename for a working copy of a template."""
        base, ext = os.path.splitext(filename)
        uid = uuid.uuid4().hex[:8]
        return f"{base}_copy_{uid}{ext}"

    def _replace_in_paragraph(self, paragraph, placeholder: str, value: str) -> int:
        """Replace placeholder in one paragraph, return number of replacements."""
        full_text = self._para_text(paragraph)
        if placeholder not in full_text:
            return 0

        count = full_text.count(placeholder)
        if count == 0:
            return 0

        new_text = full_text.replace(placeholder, value)
        self._set_para_text(paragraph, new_text)
        return count

    def _replace_placeholder_in_doc(
        self, doc: Document, placeholder: str, value: str
    ) -> int:
        """Replace placeholder in all paragraphs and table cells."""
        total = 0

        for paragraph in doc.paragraphs:
            total += self._replace_in_paragraph(paragraph, placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total += self._replace_in_paragraph(
                            paragraph, placeholder, value
                        )

        return total

    def _fill_field_by_label(self, doc: Document, field_label: str, value: str) -> bool:
        """Fill a value after a given label (paragraph or table cell)."""
        label_norm = field_label.strip()
        if not label_norm:
            return False

        def matches(text: str) -> bool:
            t = text.strip()
            if not t:
                return False
            low = t.lower()
            if label_norm.lower() in low:
                return True
            if (label_norm + ":").lower() in low:
                return True
            return False

        updated = False

        # Paragraph labels
        paragraphs = doc.paragraphs
        for i, p in enumerate(paragraphs):
            txt = self._para_text(p)
            if matches(txt):
                if i + 1 < len(paragraphs):
                    next_txt = self._para_text(paragraphs[i + 1]).strip()
                    if not next_txt:
                        self._set_para_text(paragraphs[i + 1], value)
                    else:
                        base = txt.rstrip()
                        sep = " " if base.endswith(":") else ": "
                        self._set_para_text(p, base + sep + value)
                else:
                    base = txt.rstrip()
                    sep = " " if base.endswith(":") else ": "
                    self._set_para_text(p, base + sep + value)
                updated = True

        # Table labels
        for table in doc.tables:
            for row in table.rows:
                for c_idx, cell in enumerate(row.cells):
                    cell_txt = self._cell_text(cell)
                    if matches(cell_txt):
                        if c_idx + 1 < len(row.cells):
                            target_cell = row.cells[c_idx + 1]
                            current = self._cell_text(target_cell)
                            self._set_cell_text(
                                target_cell,
                                value if not current else f"{current} {value}",
                            )
                        else:
                            base = cell_txt.rstrip()
                            sep = " " if base.endswith(":") else ": "
                            self._set_cell_text(cell, base + sep + value)
                        updated = True

        return updated

    def _select_option(self, doc: Document, field_label: str, option_text: str) -> bool:
        """
        In a 'select an option' block (e.g. under 'Způsob dopravy'),
        keep only the selected option and clear the others.
        Selected option is visually marked by replacing '☐' with '☒'.
        """
        label_norm = field_label.strip().lower()
        option_norm = option_text.strip().lower()
        if not label_norm or not option_norm:
            return False

        updated = False
        paragraphs = doc.paragraphs

        for i, p in enumerate(paragraphs):
            txt = self._para_text(p)
            if label_norm in txt.lower():
                # We found the "group" label. Now scan following paragraphs for options.
                options_started = False
                for j in range(i + 1, len(paragraphs)):
                    opt_par = paragraphs[j]
                    opt_txt = self._para_text(opt_par)
                    stripped = opt_txt.strip()

                    # Stop when we hit a blank line after options started
                    if not stripped and options_started:
                        break

                    # Look for lines with checkbox-like markers
                    if "☐" in stripped or "[ ]" in stripped:
                        options_started = True
                        # Decide if this paragraph is the selected one
                        if option_norm in stripped.lower():
                            # Mark as selected
                            new_txt = stripped.replace("☐", "☒", 1)
                            new_txt = new_txt.replace("[ ]", "[X]", 1)
                            self._set_para_text(opt_par, new_txt)
                            updated = True
                        else:
                            # Clear other options
                            self._set_para_text(opt_par, "")
                    else:
                        # If we've already passed some options and this line isn't one, end block
                        if options_started:
                            break
                break

        return updated
        return updated

    def _find_best_matching_doc(self, query: str) -> tuple[str | None, str]:
        """Find the best matching .docx template name for the query."""
        files = self._list_docx_files()
        if not files:
            return None, "V adresáři šablon nejsou žádné .docx soubory."

        q = query.strip().lower()
        if not q:
            return None, "Dotaz je prázdný."

        for f in files:
            if f.lower() == q:
                return f, f"Použit přesný název souboru '{f}'."

        contains = [f for f in files if q in f.lower()]
        if len(contains) == 1:
            return contains[0], f"Nalezeno podle podřetězce v názvu souboru '{contains[0]}'."
        if len(contains) > 1:
            main = contains[0]
            alts = ", ".join(contains[1:])
            return main, (
                f"Nalezeno více souborů podle popisu. Vybrán '{main}'. "
                f"Další možné: {alts}"
            )

        matches = difflib.get_close_matches(q, files, n=3, cutoff=0.3)
        if not matches:
            available = ", ".join(files)
            return None, (
                "Nepodařilo se najít vhodný dokument podle popisu. "
                f"Dostupné šablony: {available}"
            )

        main = matches[0]
        if len(matches) == 1:
            return main, f"Dokument přibližně odpovídá popisu: '{main}'."
        alts = ", ".join(matches[1:])
        return main, (
            f"Dokument nejlépe odpovídající popisu: '{main}'. "
            f"Další kandidáti: {alts}"
        )
    def create_table(self):
        create_table_query = f"""
        CREATE TABLE IF NOT EXISTS {self.table_name} (
        id INTEGER,
        filename LONGVARCHAR,
        content LONGVARCHAR,
        vector VECTOR(DOUBLE, 384)
        )
        """
        self.cursor.execute(create_table_query)
        # self.conn.commit()
        # self.conn.close()
    def vectorize_content(self, df):
        embeddings = self.get_embedding_model().encode(df['content'], normalize_embeddings=True, show_progress_bar=True)
        return embeddings
    def vectorize_filename(self, df):
        embeddings = self.get_embedding_model().encode(df['filename'], normalize_embeddings=True, show_progress_bar=True)
        return embeddings
    def insert_chunks_into_table(self, chunks: pd.DataFrame):
        chunks = pd.DataFrame(chunks)
        self.create_table()
        embeddings = self.vectorize_content(chunks)
        chunks["vector"] = embeddings.tolist()
        # IMPORTANT: convert vector (list) → string
        chunks["vector"] = chunks["vector"].apply(lambda v: str(v))
        insert_query = f"""
        INSERT INTO {self.table_name} (id, filename, content, vector)
        VALUES (?, ?, ?, TO_VECTOR(?))
        """
        rows_list = chunks[["id", "filename", "content", "vector"]].values.tolist()
        self.cursor.executemany(insert_query, rows_list)
        # self.conn.commit()
        # self.conn.close()
        print("Insertions done!")

    def _find_uploads_dir(self) -> str:
        """
        Locate an existing 'uploads' directory anywhere under the current
        working directory. If none is found, create one under doc_root.
        """
        base_dir = os.getcwd()
        for root, dirs, _ in os.walk(base_dir):
            for d in dirs:
                if d.lower() == "uploads":
                    return os.path.join(root, d)

        default_dir = os.path.join(self.doc_root, "uploads")
        os.makedirs(default_dir, exist_ok=True)
        return default_dir

    # ------------- AGENT / TOOLS ------------- #

    def create_chatbot(self):
        checkpointer = InMemorySaver()

        llm = ChatOpenAI(
            model="gpt-5.1",
            temperature=0.1,
            api_key=OPENAI_API_KEY,
        )

        @tool
        def load_word_document(query: str) -> str:
            """Load a template by name/description and create a working copy."""
            filename, info = self._find_best_matching_doc(query)
            if not filename:
                return info

            template_path = os.path.join(self.doc_root, filename)
            try:
                template_doc = Document(template_path)
            except Exception as e:
                return f"Dokument '{filename}' se nepodařilo načíst: {e}"

            working_filename = self._make_working_filename(filename)
            working_path = os.path.join(self.doc_root, working_filename)
            template_doc.save(working_path)

            doc = Document(working_path)
            self.current_doc = doc
            self.current_doc_path = working_path
            self.current_doc_name = working_filename

            preview = self._doc_to_text(doc, max_chars=1500)
            return (
                f"{info}\n\n"
                f"Původní šablona '{filename}' byla zkopírována do souboru "
                f"'{working_filename}', který je nyní připraven k vyplňování.\n\n"
                f"Náhled obsahu:\n{preview}"
            )

        @tool
        def show_current_document() -> str:
            """Return the text view of the currently loaded document."""
            if self.current_doc is None:
                return "Momentálně není načten žádný dokument. Použij nejdřív 'load_word_document'."
            text = self._doc_to_text(self.current_doc, max_chars=4000)
            return (
                f"Aktuálně načtený dokument: {self.current_doc_name}\n\n"
                f"Text dokumentu:\n{text}"
            )

        @tool
        def save_document_as(new_filename: str) -> str:
            """Save the current working document under a new .docx name."""
            if self.current_doc is None:
                return "Není načten žádný dokument, není co ukládat."

            if not new_filename.lower().endswith(".docx"):
                new_filename += ".docx"

            new_path = os.path.join(self.doc_root, new_filename)
            self.current_doc.save(new_path)

            # >>> This was missing – make the newly saved file the current one <<<
            self.current_doc_path = new_path
            self.current_doc_name = new_filename

            return (
                f"Aktualizovaný dokument byl uložen jako '{new_filename}'. "
                f"Cesta: {new_path}"
            )

        @tool
        def fill_placeholder(field_name: str, value: str) -> str:
            """
            Fill a field in the current document:
            - {{FIELD_NAME}} placeholder, or
            - Czech label like 'Jméno a příjmení:' / table label.
            """
            if self.current_doc is None or self.current_doc_path is None:
                return "Není načten žádný dokument. Použij nejdřív 'load_word_document'."

            placeholder = f"{{{{{field_name}}}}}"
            replaced_count = self._replace_placeholder_in_doc(
                self.current_doc, placeholder, value
            )

            if replaced_count == 0:
                filled_by_label = self._fill_field_by_label(
                    self.current_doc, field_name, value
                )
                if not filled_by_label:
                    return (
                        f"Nepodařilo se najít ani placeholder '{placeholder}', "
                        f"ani pole se štítkem odpovídajícím '{field_name}'. "
                        "Zkontroluj prosím název pole."
                    )

            self.current_doc.save(self.current_doc_path)
            return (
                f"Pole '{field_name}' bylo vyplněno hodnotou '{value}'. "
                f"Aktualizovaný dokument je uložen jako {self.current_doc_name}."
            )


        @tool
        def choose_option(field_label: str, option_text: str) -> str:
            """Select one checkbox option in the given section and clear others."""
            if self.current_doc is None or self.current_doc_path is None:
                return "Není načten žádný dokument. Použij nejdřív 'load_word_document'."

            changed = self._select_option(self.current_doc, field_label, option_text)
            if not changed:
                return (
                    f"Nepodařilo se najít sekci '{field_label}' nebo možnost obsahující "
                    f"'{option_text}'. Zkontroluj prosím texty v šabloně."
                )

            self.current_doc.save(self.current_doc_path)
            return (
                f"V sekci '{field_label}' byla vybrána možnost '{option_text}'. "
                f"Ostatní možnosti byly odstraněny nebo vyprázdněny. "
                f"Aktualizovaný dokument je uložen jako {self.current_doc_name}."
            )

        @tool
        def send_to_uploads() -> str:
            """
            Move the current working document into an 'uploads' folder so that
            it can be picked up by other systems.
            """
            if self.current_doc is None or self.current_doc_path is None:
                return "Není načten žádný dokument. Použij nejdřív 'load_word_document'."

            if not os.path.isfile(self.current_doc_path):
                return f"Soubor '{self.current_doc_path}' neexistuje na disku."

            uploads_dir = self._find_uploads_dir()
            filename = os.path.basename(self.current_doc_path)
            dest_path = os.path.join(uploads_dir, filename)

            if os.path.exists(dest_path):
                base, ext = os.path.splitext(filename)
                suffix = 1
                while True:
                    candidate = f"{base}_{suffix}{ext}"
                    candidate_path = os.path.join(uploads_dir, candidate)
                    if not os.path.exists(candidate_path):
                        dest_path = candidate_path
                        break
                    suffix += 1

            shutil.move(self.current_doc_path, dest_path)
            self.current_doc_path = dest_path

            return (
                "Dokument byl přesunut do složky 'uploads'. "
                f"Nová cesta: {dest_path}"
            )

        tools = [
            load_word_document,
            show_current_document,
            fill_placeholder,
            choose_option,
            save_document_as,
            send_to_uploads,
        ]

        agent = create_agent(
            model=llm,
            tools=tools,
            middleware=[
                SummarizationMiddleware(
                    model=llm,
                    trigger=("tokens", 4000),
                    keep=("messages", 20),
                ),
            ],
            checkpointer=checkpointer,
        )
        return agent

    # ------------- CHAT LOOP ------------- #

    def get_prompt(self):
        return input()

    def validation(self, result):
        return result

    def return_response(self, query):
        results = self.vector_search(query)

        system_prompt = """ Základy: 1. Jsi užitečný asistent, chatbot fungující v nemocnici. 2. Tvým posláním je odpovídat na dotazy zaměstnanců týkající se jejich práce a provádět je organizační strukturou nemocnice a administrativními procesy. 3. Poskytuj odpovědi přesně podle interních dokumentů, které jsou dostupné prostřednictvím RAG (retrieved context). 4. Uživatel má být bezpečně a krok za krokem proveden procesem či postupem tak, aby splnil veškeré požadavky směrnic a nic nevynechal. Tvůj způsob práce: 1. Odpovídej v jazyku, jakým mluví uživatel. 2. Ptej se uživatele na jeden konkrétní krok procesu. Nikdy nepřeskakuj více kroků najednou. 3. Vysvětluj pouze to, co uživatel potřebuje vědět pro aktuální krok. 4. Pokud je dotaz faktický, vždy nejprve vyhledej informace v dokumentech RAG. 5. Neodpovídej věci, které nejsou v podkladech, raději uveď, že nejsou uvedeny, nebo navrhni, kde se hledají. 6. Pokud uživatel neví, co má dělat, navrhni další krok. 7. Vyhýbej se nepodloženému nebo podlézavému lichocení. 8. Zachovej profesionalitu a střízlivou upřímnost. Co nesmíš dělat: 1. Nevymýšlej si pravidla, která nejsou ve zdrojových dokumentech. 2. Nevytvářej interní postupy, pokud nejsou výslovně uvedené. 3. Nehádej hodnoty (např. sazby stravného). Práce s dokumenty: - Pokud chce uživatel vyplnit formulář/dokument: 1. Rozhodni se, který z dostupných dokumentů a formulářů potřebuje. 2. Zavolej nástroj 'load_word_document' a jako argument použij: - buď přesný název souboru (např. 'Formular_XY.docx'), - nebo slovní popis (např. 'žádost o dovolenou', 'stížnost na dokumentaci'). 3. Pokud potřebuješ znát strukturu, použij 'show_current_document'. 4. U každé kategorie údajů (např. údaje o cestě či způsob dopravy) si vyžádej údaje o všech podúdajích od uživatele a použij nástroj 'fill_placeholder' s názvem pole nebo textovým štítkem (bez složených závorek) a hodnotou. 5. Pokud šablona obsahuje zástupné texty ve tvaru {{NAZEV_POLE}}, předávej do 'fill_placeholder' právě tento název pole. 6. Pokud formulář obsahuje pouze textové štítky jako 'Jméno a příjmení:' nebo 'Datum a čas odjezdu:', předávej tyto štítky (ideálně včetně dvojtečky) jako argument 'field_name' do nástroje 'fill_placeholder' - nástroj se pokusí doplnit hodnotu do řádku pod nebo do buňky vpravo (např. v tabulce 'Odhadované náklady'). 7. Pokud je v šabloně sekce se seznamem voleb (např. 'Způsob dopravy' s několika checkboxy), použij nástroj 'choose_option' s názvem sekce (např. 'Způsob dopravy') a textem vybrané možnosti (např. 'Soukromé vozidlo'). Nástroj nechá jen zvolenou možnost a ostatní odstraní. 8. Po dokončení použij 'save_document_as' a pojmenuj soubor podle kontextu. Originální šablona se nesmí přepsat. 9. Pokud uživatel upraví nějaké údaje, vymaž předchozí údaje a nahraď je novými. Pokud se uživatel dotazuje na nějaký proces v nemocnici (např. "Chci si stěžovat na nedostatečnou dokumentaci k webové aplikaci vyvinuté v Centru Informatiky (CI)"), 1. Odpovídej jasně a požádej uživatele o upřesnění, pokud nemůžeš přesně určit proces, který je pro uživatele relevantní (v tomto případě proces dokumentace ze strany oddělení nezdravotnických aplikací, které je součástí CI). 2. Pokud má uživatel podle předpisů více možností, jak dosáhnout svého cíle, popiš dostupné možnosti a zeptej se uživatele, kterou si chce vybrat。 - Pokud musí kontaktovat jiného zaměstnance, ale nemáš jeho kontaktní údaje, jasně mu sděl, že je nemáš. - Pokud musí kontaktovat jiného zaměstnance a ty máš jeho kontaktní údaje, poskytni mu tyto informace (jméno, telefonní číslo, e-mail). 3. Při odpovídání vždy upřednostňuj organizační informace z dodaných dokumentů. Pokud tam informace není dostupná, informuj o tom uživatele a nic si nevymýšlej. 4. Pokud nemáš informace o uživatelově dotazu nebo o tom, jak by měl uživatel v daném procesu postupovat, ale máš informace o tom, kde může uživatel získat kvalifikovanou pomoc, doporuč mu osoby, které má kontaktovat, a poskytni kontaktní informace (v tomto případě by měl uživatel kontaktovat oddělení nezdravotnických aplikací). 5. Na konci své odpovědi odkazuj k dokumentům (text "Text z dokumentu XYZ.docx", před ->), ze kterých jsi čerpal informace, pokud jsou relevantní. Vypiš je na konci odpovědi ve formátu: "Dále se můžete obrátit na dokument XYZ". 6. Pokud uživatel poprosí o pomoc s procesem, proveď ho několika kroky, které musí podniknout, aby dosáhl svého cíle. """

        context_msg = ""
        if results:
            joined = "\n".join(results)
            context_msg = (
                "Následuje kontext z nemocničních dokumentů. Při odpovědi se drž těchto informací a na konci odpovědi odkaž na příslušné dokumenty.\n\n"
                f"{joined}"
            )

        messages = [("system", system_prompt)]
        if context_msg:
            messages.append(("system", context_msg))
        messages.append(("user", query))

        response = self.agent.invoke({"messages": messages}, self.config)
        validated_response = self.validation(response)

        return validated_response["messages"][-1].content
