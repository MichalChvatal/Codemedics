"""Agent classes rewritten to use the modern OpenAI Python SDK (>=1.0.0).

Key changes:
- Removes legacy ChatCompletion.create usage.
- Uses the `OpenAI` client and `client.chat.completions.create`.
- Clean, minimal, predictable behavior.
"""

from typing import Optional, List, Dict, Any
from openai import OpenAI
import os
import difflib
import os
import uuid
from langchain.agents import create_agent
from langchain.agents.middleware import SummarizationMiddleware
from langgraph.checkpoint.memory import InMemorySaver
from openai import OpenAI

from langchain_openai import ChatOpenAI
from langchain.agents import create_agent
from langchain.agents.middleware import SummarizationMiddleware
from langgraph.checkpoint.memory    import InMemorySaver
from langchain_core.tools import tool
from docx import Document

from docx.oxml import OxmlElement   

class BaseAgent:
    def __init__(self, name: str, system_prompt: str,llm_client: Optional[OpenAI] = None):
        """Base agent wrapper.

        Args:
            name: logical name of the agent
            system_prompt: prompt to set as the system role
            llm_client: Optional OpenAI SDK client. If None, a default client is created.
        """
        self.name = name
        self.system_prompt = system_prompt
        self.tools = []
        
        self.llm_client = OpenAI(api_key="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")

    def add_tools(self, tools: List[Any]) -> None:
        """Add tools to the agent.

        Args:
            tools: list of tool functions
        """
        self.tools.extend(tools)

    def _prepare_messages(self, user_query: str, rag_context: str) -> List[Dict[str, str]]:
        messages = [{"role": "system", "content": self.system_prompt}]
        if rag_context:
            messages.append({"role": "assistant", "content": f"Relevant documents:\n{rag_context}"})
        messages.append({"role": "user", "content": user_query})
        return messages

    def invoke(self, user_query: str, rag_context: str = "") -> str:
        """Call the model using the modern OpenAI client API.
        """
        messages = self._prepare_messages(user_query, rag_context)

        try:
            response = self.llm_client.chat.completions.create(
                model="gpt-5.1",
                tools = self.tools,
                messages=messages,
                temperature=0.1,
                max_completion_tokens=1500,
            )
        except Exception as e:
            return f"[OpenAI invocation failed] {e}"

        try:
            return response.choices[0].message.content.strip()
        except Exception:
            return str(response)


FORM_AGENT_PROMPT = """
Kdo jsi:
- Jsi FormAgent, expert na šablony a formuláře nemocnice.
Tvůj úkol:
- Pomáháš uživateli krok za krokem vyplnit konkrétní formulář.
- Uživatel vyplní formulář správně, podle RAG dokumentace a struktury šablony.
- Nikdy nepřeskakuješ více kroků najednou.
Jak pracuješ:
- Identifikuj, jaký formulář uživatel potřebuje (– podle RAG kontextu nebo popisu).
- Pokud je nejasný výběr dokumentu, zeptej se na jediné upřesnění.
U každého kroku:
- vysvětli, jaké pole je potřeba vyplnit,
- vyžádej si hodnotu,
- použij nástroj fill_placeholder (nebo ekvivalent v orchestrátoru).
- Pokud šablona obsahuje:
    - {{POLE}} → předávej přesně název POLE,
    - textové štítky jako „Jméno zaměstnance:“ → předávej je jako field_name.
- Po dokončení nabídni uživateli možnost dokument uložit.
Pravidla:
- Nepřeskakuj žádné části formuláře.
- Nepřidávej pole, která nejsou v šabloně.
- Nehádej hodnoty — vždy si je vyžádej od uživatele.
- Vše vycházej výhradně z dokumentů RAG."""

PROCESS_AGENT_PROMPT = """
Kdo jsi:
- Jsi ProcessAgent, specialista na nemocniční procesy, workflow a administrativní postupy.
Úkol:
- Provést uživatele detailně a bezpečně krok za krokem procesem, který potřebuje vykonat.
- Například: služební cesta, reklamace, stížnost, žádost o přístup, nástup/ukončení pracovního poměru, evidence práce atd.
Jak pracuješ:
- Identifikuj proces podle dotazu a RAG informací.
- Pokud není jasné, který proces uživatel chce, požádej ho o jediné upřesnění.
Pracuj krokově:
- vždy popiš pouze jeden krok,
- vysvětli, co uživatel musí udělat,
- nabídni pokračování.
- Pokud existuje více možných postupů, popiš všechny a zeptej se, který chce uživatel použít.
- Pokud musí kontaktovat jiného zaměstnance a dokumenty obsahují kontakt → poskytni ho.
- Pokud dokumenty kontakt neobsahují → řekni to otevřeně.
Co nesmíš:
- Nevymýšlej procesy, které nejsou v RAG.
- Nehádej parametry (např. sazby, výše náhrad, termíny).
- Neprováděj vyplňování formuláře (předej FORM_AGENTOVI).
Cíl:
- Uživatel bezpečně projde celý proces a nic nevynechá."""

ORG_AGENT_PROMPT= """
Kdo jsi:
- Jsi OrgAgent, expert na organizační strukturu nemocnice a kontaktní místa.
Úkol:
- Pomáháš uživateli najít správné oddělení, kancelář, pracoviště nebo osobu, která má jeho problém řešit.
Jak pracuješ:
- Používej pouze informace z RAG zdrojů.
- Pokud uživatel popíše situaci:
    - identifikuj, které oddělení se tím zabývá,
    - vysvětli proč,
    - nabídni kontaktní údaje (pokud jsou v RAG).
    - Pokud je možností více, nabídni všechny.
    - Pokud dokument neobsahuje kontakt, jasně to oznam.
Co nesmíš:
- Nevymýšlej pozice, oddělení nebo kontakty.
- Neprováděj procesní kroky ani nevyplňuj formuláře.
Cíl:
- Uživatel zjistí, kam se obrátit, a dostane přesné kontaktní instrukce."""

class ProcessAgent(BaseAgent):
    def __init__(self):
        super().__init__("PROCESS_AGENT", PROCESS_AGENT_PROMPT, [])


class OrgAgent(BaseAgent):
    def __init__(self):
        super().__init__("ORG_AGENT", ORG_AGENT_PROMPT, [])

class FormAgent(BaseAgent):
    def __init__(self):
        super().__init__("FORM_AGENT", FORM_AGENT_PROMPT)
        tools = [
            self.load_word_document,
            self.show_current_document,
            self.fill_placeholder,
            self.choose_option,
            self.save_document_as,
        ]
        
        self.add_tools(tools)

# ------------- WORD DOCUMENT HELPERS ------------- #
    @staticmethod
    def _para_text(paragraph) -> str:
        """Read paragraph text via run XML (avoids lxml XPath bug)."""
        pieces: list[str] = []
        for run in paragraph.runs:
            r = run._r
            for child in r:
                tag = child.tag
                if tag.endswith("}t") and child.text:
                    pieces.append(child.text)
                elif tag.endswith("}tab"):
                    pieces.append("\t")
                elif tag.endswith(("}br", "}cr")):
                    pieces.append("\n")
        return "".join(pieces)

    def _set_para_text(self, paragraph, new_text: str) -> None:
        """Replace paragraph content with a single run containing new_text."""
        p = paragraph._p
        for child in list(p):
            p.remove(child)

        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_text
        r.append(t)
        p.append(r)

    def _cell_text(self, cell) -> str:
        """Concatenate visible text of all paragraphs in a cell."""
        return "\n".join(self._para_text(p) for p in cell.paragraphs).strip()

    def _set_cell_text(self, cell, new_text: str) -> None:
        """Overwrite a cell with a single paragraph containing new_text."""
        if not cell.paragraphs:
            p = cell.add_paragraph("")
            self._set_para_text(p, new_text)
            return

        for i, p in enumerate(cell.paragraphs):
            if i == 0:
                self._set_para_text(p, new_text)
            else:
                self._set_para_text(p, "")

    def _doc_to_text(self, doc: Document, max_chars: int = 4000) -> str:
        """Plain text view of a document, truncated if needed."""
        parts = [self._para_text(p).strip() for p in doc.paragraphs]
        parts = [p for p in parts if p]
        text = "\n".join(parts)
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[Text zkrácen kvůli délce…]"
        return text

    def _list_docx_files(self):
        if not os.path.isdir(self.doc_root):
            return []
        return [
            f
            for f in os.listdir(self.doc_root)
            if f.lower().endswith(".docx")
            and os.path.isfile(os.path.join(self.doc_root, f))
        ]

    @staticmethod
    def _make_working_filename(filename: str) -> str:
        """Generate unique filename for a working copy of a template."""
        base, ext = os.path.splitext(filename)
        uid = uuid.uuid4().hex[:8]
        return f"{base}_copy_{uid}{ext}"

    def _replace_in_paragraph(self, paragraph, placeholder: str, value: str) -> int:
        """Replace placeholder in one paragraph, return number of replacements."""
        full_text = self._para_text(paragraph)
        if placeholder not in full_text:
            return 0

        count = full_text.count(placeholder)
        if count == 0:
            return 0

        new_text = full_text.replace(placeholder, value)
        self._set_para_text(paragraph, new_text)
        return count

    def _replace_placeholder_in_doc(
        self, doc: Document, placeholder: str, value: str
    ) -> int:
        """Replace placeholder in all paragraphs and table cells."""
        total = 0

        for paragraph in doc.paragraphs:
            total += self._replace_in_paragraph(paragraph, placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total += self._replace_in_paragraph(
                            paragraph, placeholder, value
                        )

        return total

    def _fill_field_by_label(self, doc: Document, field_label: str, value: str) -> bool:
        """Fill a value after a given label (paragraph or table cell)."""
        label_norm = field_label.strip()
        if not label_norm:
            return False

        def matches(text: str) -> bool:
            t = text.strip()
            if not t:
                return False
            low = t.lower()
            if label_norm.lower() in low:
                return True
            if (label_norm + ":").lower() in low:
                return True
            return False

        updated = False

        # Paragraph labels
        paragraphs = doc.paragraphs
        for i, p in enumerate(paragraphs):
            txt = self._para_text(p)
            if matches(txt):
                if i + 1 < len(paragraphs):
                    next_txt = self._para_text(paragraphs[i + 1]).strip()
                    if not next_txt:
                        self._set_para_text(paragraphs[i + 1], value)
                    else:
                        # Append value to label paragraph
                        base = txt.rstrip()
                        sep = " " if base.endswith(":") else ": "
                        self._set_para_text(p, base + sep + value)
                else:
                    base = txt.rstrip()
                    sep = " " if base.endswith(":") else ": "
                    self._set_para_text(p, base + sep + value)
                updated = True

        # Table labels
        for table in doc.tables:
            for row in table.rows:
                for c_idx, cell in enumerate(row.cells):
                    cell_txt = self._cell_text(cell)
                    if matches(cell_txt):
                        if c_idx + 1 < len(row.cells):
                            target_cell = row.cells[c_idx + 1]
                            current = self._cell_text(target_cell)
                            self._set_cell_text(
                                target_cell,
                                value if not current else f"{current} {value}",
                            )
                        else:
                            base = cell_txt.rstrip()
                            sep = " " if base.endswith(":") else ": "
                            self._set_cell_text(cell, base + sep + value)
                        updated = True

        return updated

    def _select_option(self, doc: Document, field_label: str, option_text: str) -> bool:
        """
        In a checkbox group (e.g. 'Způsob dopravy'), keep only the chosen option
        and mark it as selected.
        """
        label_norm = field_label.strip().lower()
        option_norm = option_text.strip().lower()
        if not label_norm or not option_norm:
            return False

        updated = False
        paragraphs = doc.paragraphs

        for i, p in enumerate(paragraphs):
            group_txt = self._para_text(p)
            if label_norm in group_txt.lower():
                options_started = False
                for j in range(i + 1, len(paragraphs)):
                    opt_par = paragraphs[j]
                    opt_txt = self._para_text(opt_par)
                    stripped = opt_txt.strip()

                    if not stripped and options_started:
                        break

                    if "☐" in stripped or "[ ]" in stripped:
                        options_started = True
                        if option_norm in stripped.lower():
                            # mark selected
                            new_txt = stripped.replace("☐", "☒", 1)
                            new_txt = new_txt.replace("[ ]", "[X]", 1)
                            self._set_para_text(opt_par, new_txt)
                            updated = True
                        else:
                            # clear unselected options
                            self._set_para_text(opt_par, "")
                    else:
                        if options_started:
                            break
                break

        return updated

    def _find_best_matching_doc(self, query: str) -> tuple[str | None, str]:
        """Find the best matching .docx template name for the query."""
        files = self._list_docx_files()
        if not files:
            return None, "V adresáři šablon nejsou žádné .docx soubory."

        q = query.strip().lower()
        if not q:
            return None, "Dotaz je prázdný."

        # Exact match
        for f in files:
            if f.lower() == q:
                return f, f"Použit přesný název souboru '{f}'."

        # Substring match
        contains = [f for f in files if q in f.lower()]
        if len(contains) == 1:
            return contains[0], f"Nalezeno podle podřetězce v názvu souboru '{contains[0]}'."
        if len(contains) > 1:
            main = contains[0]
            alts = ", ".join(contains[1:])
            return main, (
                f"Nalezeno více souborů podle popisu. Vybrán '{main}'. "
                f"Další možné: {alts}"
            )

        # Fuzzy match
        matches = difflib.get_close_matches(q, files, n=3, cutoff=0.3)
        if not matches:
            available = ", ".join(files)
            return None, (
                "Nepodařilo se najít vhodný dokument podle popisu. "
                f"Dostupné šablony: {available}"
            )

        main = matches[0]
        if len(matches) == 1:
            return main, f"Dokument přibližně odpovídá popisu: '{main}'."
        alts = ", ".join(matches[1:])
        return main, (
            f"Dokument nejlépe odpovídající popisu: '{main}'. "
            f"Další kandidáti: {alts}"
        )

    # ------------- AGENT / TOOLS ------------- #

    @tool
    def load_word_document(query: str) -> str:
        """Load a template by name/description and create a working copy."""
        filename, info = self._find_best_matching_doc(query)
        if not filename:
            return info

        template_path = os.path.join(self.doc_root, filename)
        try:
            template_doc = Document(template_path)
        except Exception as e:
            return f"Dokument '{filename}' se nepodařilo načíst: {e}"

        working_filename = self._make_working_filename(filename)
        working_path = os.path.join(self.doc_root, working_filename)
        template_doc.save(working_path)

        doc = Document(working_path)
        self.current_doc = doc
        self.current_doc_path = working_path
        self.current_doc_name = working_filename

        preview = self._doc_to_text(doc, max_chars=1500)
        return (
            f"{info}\n\n"
            f"Původní šablona '{filename}' byla zkopírována do souboru "
            f"'{working_filename}', který je nyní připraven k vyplňování.\n\n"
            f"Náhled obsahu:\n{preview}"
        )

    @tool
    def show_current_document() -> str:
        """Return the text view of the currently loaded document."""
        if self.current_doc is None:
            return "Momentálně není načten žádný dokument. Použij nejdřív 'load_word_document'."
        text = self._doc_to_text(self.current_doc, max_chars=4000)
        return (
            f"Aktuálně načtený dokument: {self.current_doc_name}\n\n"
            f"Text dokumentu:\n{text}"
        )

    @tool
    def save_document_as(new_filename: str) -> str:
        """Save the current working document under a new .docx name."""
        if self.current_doc is None:
            return "Není načten žádný dokument, není co ukládat."

        if not new_filename.lower().endswith(".docx"):
            new_filename += ".docx"

        new_path = os.path.join(self.doc_root, new_filename)
        self.current_doc.save(new_path)
        return (
            f"Aktualizovaný dokument byl uložen jako '{new_filename}'. "
            f"Cesta: {new_path}"
        )

    @tool
    def fill_placeholder(field_name: str, value: str) -> str:
        """
        Fill a field in the current document:
        - {{FIELD_NAME}} placeholder, or
        - Czech label like 'Jméno a příjmení:' / table label.
        """
        if self.current_doc is None or self.current_doc_path is None:
            return "Není načten žádný dokument. Použij nejdřív 'load_word_document'."

        placeholder = f"{{{{{field_name}}}}}"
        replaced_count = self._replace_placeholder_in_doc(
            self.current_doc, placeholder, value
        )

        if replaced_count == 0:
            filled_by_label = self._fill_field_by_label(
                self.current_doc, field_name, value
            )
            if not filled_by_label:
                return (
                    f"Nepodařilo se najít ani placeholder '{placeholder}', "
                    f"ani pole se štítkem odpovídajícím '{field_name}'. "
                    "Zkontroluj prosím název pole."
                )

        self.current_doc.save(self.current_doc_path)
        return (
            f"Pole '{field_name}' bylo vyplněno hodnotou '{value}'. "
            f"Aktualizovaný dokument je uložen jako {self.current_doc_name}."
        )

    @tool
    def choose_option(field_label: str, option_text: str) -> str:
        """Select one checkbox option in the given section and clear others."""
        if self.current_doc is None or self.current_doc_path is None:
            return "Není načten žádný dokument. Použij nejdřív 'load_word_document'."

        changed = self._select_option(self.current_doc, field_label, option_text)
        if not changed:
            return (
                f"Nepodařilo se najít sekci '{field_label}' nebo možnost obsahující "
                f"'{option_text}'. Zkontroluj prosím texty v šabloně."
            )

        self.current_doc.save(self.current_doc_path)
        return (
            f"V sekci '{field_label}' byla vybrána možnost '{option_text}'. "
            f"Ostatní možnosti byly odstraněny nebo vyprázdněny. "
            f"Aktualizovaný dokument je uložen jako {self.current_doc_name}."
        )
